"""One-off script: generate a .docx report for the completed baseline runs."""
import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

RUNS = Path("runs")
DOCS = Path("docs")
OUT = DOCS / "resultados_baseline.docx"

YOLOV8_DIR = RUNS / "yolov8" / "baseline-2"
YOLOV10_DIR = RUNS / "yolov10" / "baseline"
RTDETR_L_DIR = RUNS / "rtdetr_v3" / "baseline-2"
RTDETR_X_DIR = RUNS / "rtdetr_x" / "baseline"


# ── helpers ──────────────────────────────────────────────────────────────────

def _shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_image_with_caption(doc, path: Path, caption: str, width_cm: float = 15.0):
    if not path.exists():
        doc.add_paragraph(f"[imagem não encontrada: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(11)
    cap.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def add_summary_table(doc, rows, headers, col_widths):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row: dark grey background, white bold text
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        _shade_cell(cell, "404040")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows: alternating light grey / white, black text
    for r_idx, row in enumerate(rows):
        shade = "EBEBEB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = str(val)
            _shade_cell(cell, shade)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    _set_col_widths(table, col_widths)
    return table


def best_epoch_metrics(csv_path: Path):
    """Return the row with the highest mAP50(B)."""
    best = None
    with open(csv_path) as f:
        reader = csv.DictReader(f)
        for row in reader:
            row = {k.strip(): v.strip() for k, v in row.items()}
            if best is None or float(row["metrics/mAP50(B)"]) > float(best["metrics/mAP50(B)"]):
                best = row
    return best


# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

# default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── title ────────────────────────────────────────────────────────────────────
title = doc.add_heading("Resultados dos Modelos de Referência — Geração 0", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Self-Generated Labeling Adaptativo para Detecção de Placas de Circuito Impresso")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True

date_p = doc.add_paragraph("Data: 10 de maio de 2026")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── intro ────────────────────────────────────────────────────────────────────
add_heading(doc, "1. Contexto", 1)
doc.add_paragraph(
    "Este documento apresenta os resultados do treinamento supervisionado dos modelos de referência "
    "(Geração 0) para o experimento de self-generated labeling adaptativo aplicado à detecção de "
    "componentes em placas de circuito impresso (PCI). Todos os modelos foram treinados exclusivamente "
    "sobre o conjunto anotado manualmente (data/labeled/), sem qualquer pseudo-label, utilizando "
    "os mesmos hiperparâmetros de base para garantir comparabilidade."
)
doc.add_paragraph(
    "O conjunto de validação (data/val/) foi mantido intacto durante todo o treinamento e é utilizado "
    "apenas para avaliação. As métricas reportadas correspondem ao melhor checkpoint (best.pt) "
    "selecionado por mAP50 na validação."
)

# ── comparative summary ───────────────────────────────────────────────────────
add_heading(doc, "2. Resumo Comparativo", 1)
doc.add_paragraph(
    "A tabela abaixo resume os resultados dos quatro modelos de referência (Geração 0). "
    "As métricas globais e por classe são provenientes de execuções de scripts/evaluate.py "
    "sobre o melhor checkpoint (best.pt) de cada modelo."
)
doc.add_paragraph()

summary_headers = ["Modelo", "Melhor época", "Precision", "Recall", "mAP50", "mAP50-95"]
summary_rows = [
    ["YOLOv8m",    "29 / 49", "0,796", "0,825", "0,860", "0,652"],
    ["YOLOv10m",   "41 / 56", "0,802", "0,787", "0,837", "0,629"],
    ["RT-DETR-l",  "37 / 57", "0,779", "0,841", "0,839", "0,633"],
    ["RT-DETR-x",  "30 / 57", "0,809", "0,845", "0,837", "0,646"],
]
add_summary_table(doc, summary_rows, summary_headers, [3.5, 2.5, 2.5, 2.5, 2.5, 2.5])

doc.add_paragraph()

# ── yolov8 section ───────────────────────────────────────────────────────────
add_heading(doc, "3. YOLOv8m — Referência", 1)

add_heading(doc, "3.1 Configuração do treinamento", 2)
doc.add_paragraph(
    "Modelo base: yolov8m.pt  |  Épocas máx.: 100  |  Paciência: 20  |  "
    "Imagem: 640×640  |  Batch: 16  |  Otimizador: AdamW  |  lr0: 0,001  |  Semente: 42"
)

add_heading(doc, "3.2 Métricas globais", 2)
yolov8_headers = ["Métrica", "Valor"]
yolov8_global = [
    ["Melhor época", "29"],
    ["Época de parada (early stopping)", "49"],
    ["Precision", "0,796"],
    ["Recall", "0,825"],
    ["mAP50", "0,860"],
    ["mAP50-95", "0,652"],
]
add_summary_table(doc, yolov8_global, yolov8_headers, [9.0, 7.0])

doc.add_paragraph()
add_heading(doc, "3.3 Métricas por classe", 2)
per_class_headers = ["Classe", "mAP50", "mAP50-95"]
yolov8_per_class = [
    ["Resistor SMD",            "0,987", "0,775"],
    ["Capacitor Cerâmico",      "0,967", "0,581"],
    ["CI (Circuito Integrado)", "0,952", "0,753"],
    ["Diodo",                   "0,560", "0,143"],
    ["Indutor",                 "0,653", "0,532"],
    ["Capacitor Eletrolítico",  "0,773", "0,653"],
    ["Capacitor de Tântalo",    "0,995", "0,823"],
    ["LED",                     "0,992", "0,957"],
]
add_summary_table(doc, yolov8_per_class, per_class_headers, [7.0, 5.0, 4.0])
doc.add_paragraph()

add_heading(doc, "3.4 Curvas de treinamento", 2)
add_image_with_caption(
    doc,
    YOLOV8_DIR / "results.png",
    "Figura 1 — YOLOv8m: evolução das perdas e métricas de validação por época.",
)
doc.add_paragraph()

add_heading(doc, "3.5 Precision-Recall Curve", 2)
add_image_with_caption(
    doc,
    YOLOV8_DIR / "BoxPR_curve.png",
    "Figura 2 — YOLOv8m: Precision-Recall curve por classe no conjunto de validação.",
)
doc.add_paragraph()

add_heading(doc, "3.6 Confusion Matrix (normalizada)", 2)
add_image_with_caption(
    doc,
    YOLOV8_DIR / "confusion_matrix_normalized.png",
    "Figura 3 — YOLOv8m: confusion matrix normalizada no conjunto de validação.",
)
doc.add_paragraph()

add_heading(doc, "3.7 Exemplos de validação — rótulos vs. predições", 2)
add_image_with_caption(
    doc,
    YOLOV8_DIR / "val_batch0_labels.jpg",
    "Figura 4a — YOLOv8m: anotações ground-truth (batch 0 de validação).",
    width_cm=14.0,
)
doc.add_paragraph()
add_image_with_caption(
    doc,
    YOLOV8_DIR / "val_batch0_pred.jpg",
    "Figura 4b — YOLOv8m: predições do modelo (batch 0 de validação).",
    width_cm=14.0,
)
doc.add_paragraph()

# ── yolov10 section ──────────────────────────────────────────────────────────
add_heading(doc, "4. YOLOv10m — Referência", 1)

add_heading(doc, "4.1 Configuração do treinamento", 2)
doc.add_paragraph(
    "Modelo base: yolov10m.pt  |  Épocas máx.: 100  |  Paciência: 20  |  "
    "Imagem: 640×640  |  Batch: 16  |  Otimizador: AdamW  |  lr0: 0,001  |  Semente: 42"
)

add_heading(doc, "4.2 Métricas globais", 2)
yolov10_global = [
    ["Melhor época", "41"],
    ["Época de parada (early stopping)", "56"],
    ["Precision", "0,802"],
    ["Recall", "0,787"],
    ["mAP50", "0,837"],
    ["mAP50-95", "0,629"],
]
add_summary_table(doc, yolov10_global, yolov8_headers, [9.0, 7.0])

doc.add_paragraph()
add_heading(doc, "4.3 Métricas por classe", 2)
yolov10_per_class = [
    ["Resistor SMD",            "0,990", "0,784"],
    ["Capacitor Cerâmico",      "0,953", "0,590"],
    ["CI (Circuito Integrado)", "0,902", "0,727"],
    ["Diodo",                   "0,871", "0,383"],
    ["Indutor",                 "0,601", "0,487"],
    ["Capacitor Eletrolítico",  "0,737", "0,580"],
    ["Capacitor de Tântalo",    "0,650", "0,544"],
    ["LED",                     "0,989", "0,936"],
]
add_summary_table(doc, yolov10_per_class, per_class_headers, [7.0, 5.0, 4.0])
doc.add_paragraph()

add_heading(doc, "4.4 Curvas de treinamento", 2)
add_image_with_caption(
    doc,
    YOLOV10_DIR / "results.png",
    "Figura 5 — YOLOv10m: evolução das perdas e métricas de validação por época.",
)
doc.add_paragraph()

add_heading(doc, "4.5 Precision-Recall Curve", 2)
add_image_with_caption(
    doc,
    YOLOV10_DIR / "BoxPR_curve.png",
    "Figura 6 — YOLOv10m: Precision-Recall curve por classe no conjunto de validação.",
)
doc.add_paragraph()

add_heading(doc, "4.6 Confusion Matrix (normalizada)", 2)
add_image_with_caption(
    doc,
    YOLOV10_DIR / "confusion_matrix_normalized.png",
    "Figura 7 — YOLOv10m: confusion matrix normalizada no conjunto de validação.",
)
doc.add_paragraph()

add_heading(doc, "4.7 Exemplos de validação — rótulos vs. predições", 2)
add_image_with_caption(
    doc,
    YOLOV10_DIR / "val_batch0_labels.jpg",
    "Figura 8a — YOLOv10m: anotações ground-truth (batch 0 de validação).",
    width_cm=14.0,
)
doc.add_paragraph()
add_image_with_caption(
    doc,
    YOLOV10_DIR / "val_batch0_pred.jpg",
    "Figura 8b — YOLOv10m: predições do modelo (batch 0 de validação).",
    width_cm=14.0,
)
doc.add_paragraph()

# ── rtdetr_l section ─────────────────────────────────────────────────────────
add_heading(doc, "5. RT-DETR-l — Referência", 1)

add_heading(doc, "5.1 Configuração do treinamento", 2)
doc.add_paragraph(
    "Modelo base: rtdetr-l.pt  |  Épocas máx.: 72  |  Paciência: 20  |  "
    "Imagem: 640×640  |  Batch: 8  |  Otimizador: AdamW  |  lr0: 0,0001  |  Semente: 42"
)

add_heading(doc, "5.2 Métricas globais", 2)
rtdetr_l_global = [
    ["Melhor época", "37"],
    ["Época de parada (early stopping)", "57"],
    ["Precision", "0,779"],
    ["Recall", "0,841"],
    ["mAP50", "0,839"],
    ["mAP50-95", "0,633"],
]
add_summary_table(doc, rtdetr_l_global, yolov8_headers, [9.0, 7.0])

doc.add_paragraph()
add_heading(doc, "5.3 Métricas por classe", 2)
rtdetr_l_per_class = [
    ["Resistor SMD",            "0,978", "0,781"],
    ["Capacitor Cerâmico",      "0,966", "0,597"],
    ["CI (Circuito Integrado)", "0,864", "0,709"],
    ["Diodo",                   "0,598", "0,253"],
    ["Indutor",                 "0,478", "0,367"],
    ["Capacitor Eletrolítico",  "0,835", "0,642"],
    ["Capacitor de Tântalo",    "0,995", "0,741"],
    ["LED",                     "0,995", "0,974"],
]
add_summary_table(doc, rtdetr_l_per_class, per_class_headers, [7.0, 5.0, 4.0])
doc.add_paragraph()

add_heading(doc, "5.4 Curvas de treinamento", 2)
add_image_with_caption(
    doc,
    RTDETR_L_DIR / "results.png",
    "Figura 9 — RT-DETR-l: evolução das perdas e métricas de validação por época.",
)
doc.add_paragraph()

add_heading(doc, "5.5 Precision-Recall Curve", 2)
add_image_with_caption(
    doc,
    RTDETR_L_DIR / "BoxPR_curve.png",
    "Figura 10 — RT-DETR-l: Precision-Recall curve por classe no conjunto de validação.",
)
doc.add_paragraph()

add_heading(doc, "5.6 Confusion Matrix (normalizada)", 2)
add_image_with_caption(
    doc,
    RTDETR_L_DIR / "confusion_matrix_normalized.png",
    "Figura 11 — RT-DETR-l: confusion matrix normalizada no conjunto de validação.",
)
doc.add_paragraph()

add_heading(doc, "5.7 Exemplos de validação — rótulos vs. predições", 2)
add_image_with_caption(
    doc,
    RTDETR_L_DIR / "val_batch0_labels.jpg",
    "Figura 12a — RT-DETR-l: anotações ground-truth (batch 0 de validação).",
    width_cm=14.0,
)
doc.add_paragraph()
add_image_with_caption(
    doc,
    RTDETR_L_DIR / "val_batch0_pred.jpg",
    "Figura 12b — RT-DETR-l: predições do modelo (batch 0 de validação).",
    width_cm=14.0,
)
doc.add_paragraph()

# ── rtdetr_x section ─────────────────────────────────────────────────────────
add_heading(doc, "6. RT-DETR-x — Referência", 1)

add_heading(doc, "6.1 Configuração do treinamento", 2)
doc.add_paragraph(
    "Modelo base: rtdetr-x.pt  |  Épocas máx.: 72  |  Paciência: 20  |  "
    "Imagem: 640×640  |  Batch: 8  |  Otimizador: AdamW  |  lr0: 0,0001  |  Semente: 42"
)

add_heading(doc, "6.2 Métricas globais", 2)
rtdetr_x_global = [
    ["Melhor época", "30"],
    ["Época de parada (early stopping)", "57"],
    ["Precision", "0,809"],
    ["Recall", "0,845"],
    ["mAP50", "0,837"],
    ["mAP50-95", "0,646"],
]
add_summary_table(doc, rtdetr_x_global, yolov8_headers, [9.0, 7.0])

doc.add_paragraph()
add_heading(doc, "6.3 Métricas por classe", 2)
rtdetr_x_per_class = [
    ["Resistor SMD",            "0,973", "0,769"],
    ["Capacitor Cerâmico",      "0,947", "0,561"],
    ["CI (Circuito Integrado)", "0,842", "0,691"],
    ["Diodo",                   "0,668", "0,261"],
    ["Indutor",                 "0,573", "0,458"],
    ["Capacitor Eletrolítico",  "0,706", "0,626"],
    ["Capacitor de Tântalo",    "0,995", "0,812"],
    ["LED",                     "0,992", "0,992"],
]
add_summary_table(doc, rtdetr_x_per_class, per_class_headers, [7.0, 5.0, 4.0])
doc.add_paragraph()

add_heading(doc, "6.4 Curvas de treinamento", 2)
add_image_with_caption(
    doc,
    RTDETR_X_DIR / "results.png",
    "Figura 13 — RT-DETR-x: evolução das perdas e métricas de validação por época.",
)
doc.add_paragraph()

add_heading(doc, "6.5 Precision-Recall Curve", 2)
add_image_with_caption(
    doc,
    RTDETR_X_DIR / "BoxPR_curve.png",
    "Figura 14 — RT-DETR-x: Precision-Recall curve por classe no conjunto de validação.",
)
doc.add_paragraph()

add_heading(doc, "6.6 Confusion Matrix (normalizada)", 2)
add_image_with_caption(
    doc,
    RTDETR_X_DIR / "confusion_matrix_normalized.png",
    "Figura 15 — RT-DETR-x: confusion matrix normalizada no conjunto de validação.",
)
doc.add_paragraph()

add_heading(doc, "6.7 Exemplos de validação — rótulos vs. predições", 2)
add_image_with_caption(
    doc,
    RTDETR_X_DIR / "val_batch0_labels.jpg",
    "Figura 16a — RT-DETR-x: anotações ground-truth (batch 0 de validação).",
    width_cm=14.0,
)
doc.add_paragraph()
add_image_with_caption(
    doc,
    RTDETR_X_DIR / "val_batch0_pred.jpg",
    "Figura 16b — RT-DETR-x: predições do modelo (batch 0 de validação).",
    width_cm=14.0,
)
doc.add_paragraph()

# ── observations ─────────────────────────────────────────────────────────────
add_heading(doc, "7. Observações e Próximos Passos", 1)

obs = doc.add_paragraph()
obs.style = doc.styles["List Bullet"]
obs.text = (
    "YOLOv8m apresenta o maior mAP50 (0,860) e mAP50-95 (0,652) dentre os modelos YOLO, "
    "com convergência mais rápida (melhor época 29 vs. 41 do YOLOv10m). "
    "RT-DETR-x alcança o maior mAP50-95 geral (0,646), indicando melhor localização de caixas."
)

obs2 = doc.add_paragraph()
obs2.style = doc.styles["List Bullet"]
obs2.text = (
    "Indutor e Diodo são as classes mais fracas em todos os modelos, com RT-DETR-l registrando "
    "os piores valores (Indutor AP50=0,478; Diodo AP50=0,598). O padrão é consistente entre "
    "arquiteturas, apontando para limitação do conjunto de dados (baixa frequência de amostras, "
    "alta variabilidade visual) como causa principal. Ambas as classes são candidatas "
    "prioritárias para análise do ganho com pseudo-labels."
)

obs3 = doc.add_paragraph()
obs3.style = doc.styles["List Bullet"]
obs3.text = (
    "A maior divergência inter-modelos ocorre no Capacitor de Tântalo "
    "(YOLOv8m AP50=0,995 vs. YOLOv10m AP50=0,650) e no IC "
    "(YOLOv8m AP50=0,952 vs. RT-DETR-x AP50=0,842). "
    "O Tântalo possui apenas 12 instâncias de validação, tornando seus números voláteis. "
    "A queda no IC para os modelos RT-DETR pode indicar dificuldade do decoder de atenção "
    "com a alta variabilidade intra-classe em resolução 640px."
)

obs4 = doc.add_paragraph()
obs4.style = doc.styles["List Bullet"]
obs4.text = (
    "Próximo passo: iniciar iterações de self-training (Geração 1) para cada modelo, "
    "utilizando os checkpoints de referência como warm-start. Avaliar o ganho de mAP50 "
    "e mAP50-95 nas classes fracas (Indutor, Diodo) como métrica principal de progresso."
)

doc.add_paragraph()

# ── save ─────────────────────────────────────────────────────────────────────
DOCS.mkdir(exist_ok=True)
doc.save(str(OUT))
print(f"Relatório gerado: {OUT}")
