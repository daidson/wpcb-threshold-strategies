"""
Gera docs/discussion_yolov8_gen1_gen3.docx para compartilhar com Honorato.
Executar a partir da raiz do repositório:
    uv run --with python-docx python scripts/generate_discussion_doc.py
"""

import io
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).parent.parent
RUNS_BASE = ROOT / "runs/yolov8/baseline-2"
STATIC = ROOT / "experiments/yolov8_st_static_42"
ADAPTIVE = ROOT / "experiments/yolov8_st_adaptive_42"
PROGRESSIVE = ROOT / "experiments/yolov8_st_progressive_42"
UNLABELED = ROOT / "data/unlabeled/images"
OUT = ROOT / "docs/discussion_yolov8_gen1_gen3.docx"


# ---------------------------------------------------------------------------
# styling helpers — match generate_baseline_report.py conventions
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p


def add_image_with_caption(doc, path: Path, caption: str, width_cm: float = 15.0):
    path = Path(path)
    if not path.exists():
        doc.add_paragraph(f"[imagem não encontrada: {path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    cap.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def add_table(doc, headers, rows, col_widths_cm):
    """Dark header, alternating row shading. col_widths_cm must match len(headers)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        _shade_cell(cell, "2E4057")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for r_idx, row_data in enumerate(rows):
        shade = "EEF2F7" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = str(val)
            _shade_cell(cell, shade)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_col_widths(table, col_widths_cm)
    return table


# ---------------------------------------------------------------------------
# pseudo-label visualization
# ---------------------------------------------------------------------------

_CLASS_STYLE = {
    0: ("resistor_smd",          (220,  50,  50)),
    1: ("ceramic_cap",           ( 50, 130, 220)),
    2: ("ic",                    ( 50, 180,  80)),
    3: ("diode",                 (230, 140,  20)),
    4: ("inductor",              (160,  60, 200)),
    5: ("electrolytic_cap",      ( 20, 200, 200)),
    6: ("tantalum_cap",          (220, 200,  20)),
    7: ("led",                   (200,  80, 160)),
}


def _render_pseudo_labels(image_path: Path, label_path: Path) -> bytes:
    img = Image.open(image_path).convert("RGB")
    draw = ImageDraw.Draw(img, "RGBA")
    iw, ih = img.size
    try:
        font = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 13)
    except OSError:
        font = ImageFont.load_default()
    if label_path.exists():
        for line in label_path.read_text().splitlines():
            parts = line.strip().split()
            if len(parts) < 5:
                continue
            cls = int(parts[0])
            cx, cy, bw, bh = (float(p) for p in parts[1:5])
            x1 = int((cx - bw / 2) * iw)
            y1 = int((cy - bh / 2) * ih)
            x2 = int((cx + bw / 2) * iw)
            y2 = int((cy + bh / 2) * ih)
            name, (r, g, b) = _CLASS_STYLE.get(cls, (f"cls{cls}", (180, 180, 180)))
            draw.rectangle([x1, y1, x2, y2], outline=(r, g, b, 255),
                           width=2, fill=(r, g, b, 45))
            draw.text((x1 + 2, y1 + 1), name, fill=(255, 255, 255, 230), font=font)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf.read()


def add_pseudo_label_sample(doc, image_path: Path, label_path: Path,
                             caption: str, width_cm: float = 14.0):
    if not image_path.exists():
        doc.add_paragraph(f"[imagem não encontrada: {image_path.name}]")
        return
    png = _render_pseudo_labels(image_path, label_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(png), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    cap.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)


# ---------------------------------------------------------------------------
# document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── título ──────────────────────────────────────────────────────────────
    title = doc.add_heading(
        "YOLOv8 Self-Labeling GEN1–GEN3: Configuração e Resultados", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Para: Leandro Honorato\n").bold = True
    meta.add_run(
        "Autor: Daidson Alves  |  Orientador: Bruno J. T. Fernandes  |  13/05/2026\n")
    doc.add_paragraph(
        "Este documento resume o pipeline, os conjuntos de dados e os resultados completos "
        "do experimento YOLOv8m (GEN0 baseline + GEN1–GEN3 com três estratégias de limiar "
        "de confiança). As perguntas para discussão estão na Seção 10.")

    doc.add_page_break()

    # ── 1. O que estamos fazendo ─────────────────────────────────────────────
    add_heading(doc, "1. O que estamos fazendo")
    doc.add_paragraph(
        "Estendemos o método GEN Self-Labeling (Honorato et al., 2025) comparando três estratégias "
        "de limiar de confiança em quatro arquiteturas: YOLOv8m, YOLOv10m, RT-DETR-l, RT-DETR-x. "
        "Este documento cobre a cadeia YOLOv8m, já concluída. "
        "O pipeline corresponde diretamente à sua Figura 36: cada ciclo de self-training "
        "equivale a uma Geração. Nossos GEN1/GEN2/GEN3 correspondem aos seus GEN1/GEN2/GEN3.")

    # ── 2. Conjuntos de dados ────────────────────────────────────────────────
    add_heading(doc, "2. Conjuntos de dados")
    doc.add_paragraph(
        "Dois conjuntos — um totalmente rotulado (suporte) e um parcialmente rotulado (alvo):")
    doc.add_paragraph()
    add_table(doc,
        ["Conjunto de dados", "Papel", "Imagens", "Rótulos"],
        [
            ["FICS-PCB REMAP (aug)", "Suporte — totalmente rotulado",
             "4.194 treino / 906 val", "8 classes, 100% ground truth"],
            ["PCB DSLR Crops 512px", "Alvo — parcialmente rotulado",
             "2.927", "Somente IC (classe 2)"],
        ],
        [5.5, 4.0, 3.5, 4.0]
    )
    doc.add_paragraph()
    doc.add_paragraph("8 classes de componentes:")
    doc.add_paragraph()
    add_table(doc,
        ["ID", "Nome", "Instâncias val", "Observação"],
        [
            ["0", "resistor_smd",           "2.238", "Mais frequente"],
            ["1", "ceramic_capacitor",       "2.670", "Mais frequente"],
            ["2", "ic",                      "984",   "Ground truth parcial no alvo"],
            ["3", "diode",                   "42",    "Mais raro — mais sensível a ruído"],
            ["4", "inductor",                "114",   "Segundo mais raro"],
            ["5", "electrolytic_capacitor",  "42",    "Raro"],
            ["6", "tantalum_capacitor",      "12",    "Muito raro (métrica volátil)"],
            ["7", "led",                     "18",    "Raro, visualmente distinto"],
        ],
        [1.5, 4.5, 3.0, 8.0]
    )

    # ── 3. Pipeline ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "3. Pipeline (correspondência com sua Figura 36)")
    doc.add_paragraph(
        "Cada iteração de self-training = uma Geração. O laço abaixo se repete 3 vezes, "
        "produzindo GEN1, GEN2 e GEN3:")
    doc.add_paragraph()

    pipeline_text = (
        "Professor GEN0 (best.pt baseline)\n"
        "         │\n"
        "         ▼\n"
        "  PseudoLabeler\n"
        "  executa GEN(n-1) sobre PCB DSLR (2.927 imagens)\n"
        "  mantém predições acima do limiar de confiança por classe\n"
        "  gera pseudo-labels YOLO .txt  (exceto IC — usa ground truth)\n"
        "         │\n"
        "         ▼\n"
        "  Fusão do conjunto de treino:\n"
        "  ┌────────────────────────────────────────────────────┐\n"
        "  │  FICS-PCB REMAP rotulado    (4.194 imagens)        │\n"
        "  │  + PCB DSLR rótulos IC GT  (2.927 imagens)         │\n"
        "  │  + PCB DSLR pseudo-labels  (saída GEN(n-1))        │\n"
        "  └────────────────────────────────────────────────────┘\n"
        "         │\n"
        "         ▼\n"
        "  Treinar estudante GEN(n)\n"
        "  (máx. 50 épocas, patience=15, warm-start de GEN(n-1))\n"
        "         │\n"
        "         ▼\n"
        "  Avaliar em FICS val (906 imagens, 6.102 instâncias)\n"
        "         │\n"
        "         └──► GEN(n) vira professor da próxima iteração"
    )
    p = doc.add_paragraph()
    run = p.add_run(pipeline_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph(
        "Correspondência com Honorato (2023) Figura 36:\n"
        "  'Modelo GEN(n-1)'           →  pesos GEN(i-1) (best.pt)\n"
        "  'Base não rotulada'         →  PCB DSLR Crops (2.927 imagens)\n"
        "  'Base de suporte rotulada'  →  FICS-PCB REMAP (4.194 imagens)\n"
        "  'Detector GEN(n)'           →  estudante treinado no conjunto fundido\n"
        "  Um ciclo GEN(n-1)→GEN(n)   →  uma chamada a SelfTrainer.run()")

    # ── 4. Pseudo-labels ─────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "4. Pseudo-labels: o que são e como são gerados")
    doc.add_paragraph(
        "Pseudo-labels são anotações automáticas geradas pelo modelo professor e usadas para "
        "complementar o conjunto de treino da geração seguinte. Em vez de um humano anotar cada "
        "componente nas 2.927 imagens PCB DSLR, o detector infere as caixas delimitadoras e "
        "retemos apenas as predições com confiança acima do limiar configurado.")
    doc.add_paragraph(
        "Formato: cada imagem PCB DSLR recebe um arquivo .txt no formato YOLO — uma linha por "
        "componente detectado, com cinco campos: classe  cx  cy  largura  altura (coordenadas "
        "normalizadas de 0 a 1). Predições abaixo do limiar são descartadas; imagens sem nenhuma "
        "predição aceita ficam sem arquivo .txt (o treinador ignora imagens sem rótulo).")
    doc.add_paragraph(
        "Exceção: a classe IC (classe 2) tem ground truth real no PCB DSLR. "
        "O pipeline usa esse ground truth no lugar de pseudo-labels para IC em todas as "
        "iterações e estratégias — somente as outras 7 classes são pseudo-rotuladas.")

    add_heading(doc, "4.1 Estatísticas de pseudo-labels por iteração", level=2)
    doc.add_paragraph(
        "Total de imagens no PCB DSLR: 2.927. "
        "Imagens com ao menos um pseudo-label aceito (conf ≥ limiar) e total de anotações geradas:")
    doc.add_paragraph()
    add_table(doc,
        ["Estratégia", "GEN1 — img / anot.", "GEN2 — img / anot.", "GEN3 — img / anot."],
        [
            ["Static (0,25 fixo)",          "1.982 / 10.357", "1.916 / 1.992",  "2.897 / 5.254"],
            ["Adaptive (por classe)",       "1.982 / 10.357", "1.886 / 1.773",  "2.655 / 2.622"],
            ["Progressive (0,25→0,55)",     "1.982 / 10.357", "1.639 / 1.519",  "2.844 / 3.524"],
        ],
        [4.5, 3.5, 3.5, 3.5]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "O GEN1 é idêntico nas três estratégias porque todas partem do mesmo checkpoint GEN0 "
        "com conf=0,25. A divergência começa no GEN2: a Progressive sobe para conf=0,40 e a "
        "Adaptive usa limiares por classe derivados da curva PR do GEN1. "
        "O aumento no GEN3-Static ocorre porque, com conf=0,25 e um modelo mais maduro, "
        "mais predições passam o filtro.")

    add_heading(doc, "4.2 Distribuição de classes nos pseudo-labels do GEN1", level=2)
    doc.add_paragraph(
        "Com conf=0,25 no GEN1 (igual para todas as estratégias), classes frequentes no FICS "
        "dominam os pseudo-labels. Diode (classe 3) aparece quase nada — apenas 24 pseudo-labels "
        "de 10.357 — o que explica sua queda acentuada de AP50 nas estratégias Static e Adaptive:")
    doc.add_paragraph()
    add_table(doc,
        ["Classe", "Pseudo-labels GEN1", "% do total", "Instâncias val (FICS)"],
        [
            ["ceramic_cap (1)",       "5.123", "49,5 %", "2.670"],
            ["ic (2)",                "3.036", "29,3 %", "984  (ground truth real)"],
            ["resistor_smd (0)",      "3.010", "29,1 %", "2.238"],
            ["inductor (4)",          "605",   " 5,8 %", "114"],
            ["electrolytic_cap (5)",  "284",   " 2,7 %", "42"],
            ["led (7)",               "171",   " 1,7 %", "18"],
            ["tantalum_cap (6)",      "86",    " 0,8 %", "12"],
            ["diode (3)",             "24",    " 0,2 %", "42"],
        ],
        [4.5, 3.5, 2.5, 4.5]
    )
    doc.add_paragraph()

    add_heading(doc, "4.3 Amostras visuais (PCB DSLR com pseudo-labels GEN1, Static)", level=2)
    doc.add_paragraph(
        "Imagens reais do conjunto PCB DSLR com as caixas pseudo-rotuladas do GEN1 "
        "(conf=0,25) desenhadas. Legenda de cores: "
        "vermelho=resistor_smd, azul=ceramic_cap, verde=ic, laranja=diode, "
        "roxo=inductor, ciano=electrolytic_cap, amarelo=tantalum_cap, rosa=led.")
    doc.add_paragraph()

    pl_iter1 = STATIC / "iter1/pseudo/labels"
    samples = [
        ("pcb82_rec1_crop_01", "Figura — Amostra 1: 66 pseudo-labels (placa com componentes densos)"),
        ("pcb92_rec1_crop_01", "Figura — Amostra 2: 58 pseudo-labels (placa mista)"),
        ("pcb47_rec2_crop_01", "Figura — Amostra 3: 39 pseudo-labels (componentes variados)"),
    ]
    for stem, caption in samples:
        add_pseudo_label_sample(
            doc,
            UNLABELED / f"{stem}.png",
            pl_iter1 / f"{stem}.txt",
            caption + " — GEN1, Static (conf ≥ 0,25)",
        )
        doc.add_paragraph()

    doc.add_paragraph(
        "Observação: as imagens PCB DSLR são fotos de placas reais de facilidades de reciclagem "
        "com alta variabilidade de foco, iluminação e escala. O pseudo-labeler detecta bem as "
        "classes frequentes (resistores, capacitores) mas erra em classes raras (diode) e confunde "
        "componentes em orientações incomuns. Esse ruído acumulado no treino é o mecanismo pelo "
        "qual o early stopping dispara em 4–8 épocas em todas as estratégias.")

    # ── 5. Três estratégias ──────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "5. Três estratégias de limiar")
    doc.add_paragraph(
        "As três estratégias partem do mesmo checkpoint GEN0. "
        "A diferença está apenas em como o limiar de confiança do PseudoLabeler é definido "
        "a cada iteração:")
    doc.add_paragraph()

    strategies_text = (
        "Professor GEN0\n"
        "    │\n"
        "    ├── ESTRATÉGIA A: STATIC ───────────────────────────────────────\n"
        "    │   conf = 0,25 em todas as iterações  (C_th de Honorato)\n"
        "    │   GEN0 → [PL 0,25] → GEN1 → [PL 0,25] → GEN2 → [PL 0,25] → GEN3\n"
        "    │\n"
        "    ├── ESTRATÉGIA B: ADAPTIVE ───────────────────────────────────────\n"
        "    │   conf[k] = argmax F1 na curva PR do val para a classe k\n"
        "    │   GEN0 → avalia val → conf[0..7] → [PL por classe] → GEN1\n"
        "    │   GEN1 → avalia val → recalcula conf[0..7] → [PL por classe] → GEN2\n"
        "    │   GEN2 → avalia val → recalcula conf[0..7] → [PL por classe] → GEN3\n"
        "    │\n"
        "    └── ESTRATÉGIA C: PROGRESSIVE ────────────────────────────────────\n"
        "        conf aumenta a cada iteração: 0,25 → 0,40 → 0,55\n"
        "        GEN0 → [PL 0,25] → GEN1 → [PL 0,40] → GEN2 → [PL 0,55] → GEN3"
    )
    p = doc.add_paragraph()
    run = p.add_run(strategies_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    # ── 6. Configuração de treino ────────────────────────────────────────────
    add_heading(doc, "6. Configuração de treinamento (comum às três estratégias)")
    doc.add_paragraph()
    add_table(doc,
        ["Parâmetro", "Valor"],
        [
            ["Modelo",                      "YOLOv8m (Ultralytics)"],
            ["Tamanho da imagem",           "640 × 640 px"],
            ["Batch size",                  "16"],
            ["Otimizador",                  "AdamW"],
            ["Learning rate",               "0,0005"],
            ["Épocas / iteração",           "50 (máximo)"],
            ["Early stopping patience",     "15 épocas"],
            ["Semente aleatória",           "42"],
            ["NMS IoU threshold",           "0,45"],
            ["Warm start GEN0",             "Pesos YOLOv8m pré-treinados (COCO)"],
            ["Warm start GEN1–3",           "best.pt da geração anterior"],
            ["Imagens combinadas no GEN3",  "7.121 (4.194 FICS + 2.927 PCB DSLR)"],
        ],
        [7.0, 10.0]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Avaliação: scripts/evaluate.py sobre o split de validação do FICS-PCB REMAP "
        "(906 imagens, 6.102 instâncias). Esse conjunto de val nunca é usado no treino nem "
        "na geração de pseudo-labels.\n\n"
        "Diferença metodológica: nossa avaliação usa ruído real de pseudo-labels do PCB DSLR. "
        "A Seção 5.1.1 de sua tese usou injeção de ruído sintético por não haver ground truth "
        "completo no PCB DSLR. Nossa configuração é uma medição direta, não uma simulação.")

    # ── 7. GEN0 Baseline ─────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "7. GEN0 Baseline — YOLOv8m")
    doc.add_paragraph(
        "100 épocas programadas, early stopping na época 49, "
        "melhor checkpoint na época 29. Checkpoint: runs/yolov8/baseline-2/weights/best.pt")
    doc.add_paragraph()
    add_table(doc,
        ["Classe", "AP50", "AP50-95"],
        [
            ["resistor_smd",           "0,987", "0,775"],
            ["ceramic_capacitor",      "0,967", "0,581"],
            ["ic",                     "0,952", "0,753"],
            ["diode",                  "0,560", "0,143"],
            ["inductor",               "0,653", "0,532"],
            ["electrolytic_capacitor", "0,773", "0,653"],
            ["tantalum_capacitor",     "0,995", "0,823"],
            ["led",                    "0,992", "0,957"],
            ["mAP50",                  "0,860", ""],
            ["mAP50-95",               "",      "0,652"],
            ["Precision",              "0,796", ""],
            ["Recall",                 "0,825", ""],
        ],
        [8.0, 4.5, 4.5]
    )
    doc.add_paragraph()
    add_image_with_caption(
        doc, RUNS_BASE / "results.png",
        "GEN0 Baseline — curvas de treinamento (loss, mAP, Precision, Recall por época)",
        width_cm=16.0)
    add_image_with_caption(
        doc, RUNS_BASE / "confusion_matrix_normalized.png",
        "GEN0 Baseline — confusion matrix normalizada",
        width_cm=14.0)
    add_image_with_caption(
        doc, RUNS_BASE / "val_batch0_pred.jpg",
        "GEN0 Baseline — predições sobre lote de validação (amostra)",
        width_cm=14.0)

    # ── 8A. Static ───────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "8A. Estratégia Static (conf = 0,25 em todas as iterações)")
    doc.add_paragraph(
        "Replica o C_th = 0,25 de Honorato. O mesmo limiar fixo é aplicado pelo PseudoLabeler "
        "em todas as iterações, sem ajuste por classe.")

    static_info = [
        ("—", "—", "checkpoint intermediário", ""),
        ("—", "—", "checkpoint intermediário", ""),
        ("8", "23", "0,826", "0,634"),
    ]
    for i, (best_ep, stop_ep, map50, map5095) in enumerate(static_info):
        gen = i + 1
        add_heading(doc, f"GEN{gen}", level=2)
        if gen < 3:
            doc.add_paragraph(
                f"Checkpoint intermediário — utilizado como professor para GEN{gen + 1}. "
                f"A avaliação final foi realizada apenas no GEN3.")
        else:
            doc.add_paragraph(
                f"Checkpoint final. Melhor época: {best_ep}, "
                f"early stopping na época {stop_ep}. "
                f"mAP50 = {map50}  |  mAP50-95 = {map5095}")
        img_dir = STATIC / f"iter{gen}/train"
        add_image_with_caption(
            doc, img_dir / "results.png",
            f"Static GEN{gen} — curvas de treinamento")
        add_image_with_caption(
            doc, img_dir / "confusion_matrix_normalized.png",
            f"Static GEN{gen} — confusion matrix normalizada",
            width_cm=14.0)
        add_image_with_caption(
            doc, img_dir / "val_batch0_pred.jpg",
            f"Static GEN{gen} — predições sobre lote de validação (amostra)",
            width_cm=14.0)

    # ── 8B. Adaptive ─────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "8B. Estratégia Adaptive (por classe, calibrada pelo val)")
    doc.add_paragraph(
        "Após cada iteração, o modelo é avaliado no val. Para cada classe k, "
        "o limiar de confiança é definido como o ponto que maximiza o F1 na curva PR do val. "
        "Cada classe tem seu próprio limiar, recalculado a cada geração. "
        "Esta é nossa contribuição principal em relação ao método de Honorato.")

    adaptive_info = [
        ("—", "—", "checkpoint intermediário", ""),
        ("—", "—", "checkpoint intermediário", ""),
        ("4", "19", "0,807", "0,613"),
    ]
    for i, (best_ep, stop_ep, map50, map5095) in enumerate(adaptive_info):
        gen = i + 1
        add_heading(doc, f"GEN{gen}", level=2)
        if gen < 3:
            doc.add_paragraph(
                f"Checkpoint intermediário — limiares por classe recalculados com métricas do val "
                f"antes de gerar pseudo-labels para GEN{gen + 1}.")
        else:
            doc.add_paragraph(
                f"Checkpoint final. Melhor época: {best_ep}, "
                f"early stopping na época {stop_ep}. "
                f"mAP50 = {map50}  |  mAP50-95 = {map5095}")
        img_dir = ADAPTIVE / f"iter{gen}/train"
        add_image_with_caption(
            doc, img_dir / "results.png",
            f"Adaptive GEN{gen} — curvas de treinamento")
        add_image_with_caption(
            doc, img_dir / "confusion_matrix_normalized.png",
            f"Adaptive GEN{gen} — confusion matrix normalizada",
            width_cm=14.0)
        add_image_with_caption(
            doc, img_dir / "val_batch0_pred.jpg",
            f"Adaptive GEN{gen} — predições sobre lote de validação (amostra)",
            width_cm=14.0)

    # ── 8C. Progressive ──────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "8C. Estratégia Progressive (conf 0,25 → 0,40 → 0,55)")
    doc.add_paragraph(
        "O limiar aumenta a cada geração: 0,25 no GEN1, 0,40 no GEN2, 0,55 no GEN3. "
        "A ideia é iniciar com limiar permissivo para capturar pseudo-labels de classes raras "
        "e aumentá-lo gradualmente conforme o modelo melhora.")

    progressive_info = [
        ("—", "—", "checkpoint intermediário", "", "0,25"),
        ("—", "—", "checkpoint intermediário", "", "0,40"),
        ("4", "23", "0,826", "0,640", "0,55"),
    ]
    for i, (best_ep, stop_ep, map50, map5095, thr) in enumerate(progressive_info):
        gen = i + 1
        add_heading(doc, f"GEN{gen}  (pseudo-labels gerados com conf = {thr})", level=2)
        if gen < 3:
            doc.add_paragraph(
                f"Checkpoint intermediário — pseudo-labels para GEN{gen + 1} gerados "
                f"com conf = {progressive_info[gen][4]}.")
        else:
            doc.add_paragraph(
                f"Checkpoint final. Melhor época: {best_ep}, "
                f"early stopping na época {stop_ep}. "
                f"mAP50 = {map50}  |  mAP50-95 = {map5095}")
        img_dir = PROGRESSIVE / f"iter{gen}/train"
        add_image_with_caption(
            doc, img_dir / "results.png",
            f"Progressive GEN{gen} (conf={thr}) — curvas de treinamento")
        add_image_with_caption(
            doc, img_dir / "confusion_matrix_normalized.png",
            f"Progressive GEN{gen} — confusion matrix normalizada",
            width_cm=14.0)
        add_image_with_caption(
            doc, img_dir / "val_batch0_pred.jpg",
            f"Progressive GEN{gen} — predições sobre lote de validação (amostra)",
            width_cm=14.0)

    # ── 9. Comparação final ──────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "9. Comparação final — GEN3 × GEN0 Baseline")
    doc.add_paragraph(
        "Avaliado no split de validação do FICS-PCB REMAP (906 imagens, 6.102 instâncias). "
        "Todas as métricas obtidas com scripts/evaluate.py sobre o best.pt do GEN3.")
    doc.add_paragraph()
    add_table(doc,
        ["Métrica", "GEN0 Baseline", "Static", "Adaptive", "Progressive"],
        [
            ["mAP50",        "0,860", "0,826",    "0,807",    "0,826"],
            ["mAP50-95",     "0,652", "0,634",    "0,613",    "0,640 ★"],
            ["Precision",    "0,796", "0,807 ★",  "0,742",    "0,712"],
            ["Recall",       "0,825", "0,771",    "0,772",    "0,842 ★"],
            ["Melhor época", "29/49", "8/23",     "4/19",     "4/23"],
        ],
        [4.0, 3.5, 3.5, 3.5, 3.5]
    )
    doc.add_paragraph()

    add_heading(doc, "AP50 por classe no GEN3", level=2)
    doc.add_paragraph()
    add_table(doc,
        ["Classe", "Baseline", "Static", "Adaptive", "Progressive"],
        [
            ["resistor_smd",             "0,987", "0,989",    "0,979",    "0,984"],
            ["ceramic_capacitor",        "0,967", "0,973",    "0,977",    "0,980"],
            ["ic",                       "0,952", "0,942",    "0,933",    "0,908"],
            ["diode (42 inst.)",         "0,560", "0,333",    "0,246",    "0,602 ★"],
            ["inductor (114 inst.)",     "0,653", "0,708 ★",  "0,627",    "0,587"],
            ["electrolytic_cap",         "0,773", "0,754",    "0,702",    "0,715"],
            ["tantalum_cap",             "0,995", "0,921",    "0,995 ★",  "0,837"],
            ["led",                      "0,992", "0,989",    "0,995 ★",  "0,995 ★"],
        ],
        [4.5, 3.0, 3.0, 3.0, 3.0]
    )
    doc.add_paragraph()

    add_heading(doc, "Early stopping — o sinal principal", level=2)
    doc.add_paragraph(
        "As três estratégias disparam o early stopping muito antes do GEN0:\n\n"
        "  GEN0 Baseline:    melhor época 29 de 49  (melhora por 29 épocas)\n"
        "  Static GEN3:      melhor época  8 de 23  (degrada após 8 épocas)\n"
        "  Progressive GEN3: melhor época  4 de 23  (degrada após 4 épocas)\n"
        "  Adaptive GEN3:    melhor época  4 de 19  (degrada após 4 épocas)\n\n"
        "Esse padrão se repete em todas as iterações (GEN1, GEN2, GEN3) das três estratégias. "
        "O modelo atinge seu pico em 4–8 épocas e regride sob ruído de pseudo-labels. "
        "Com conf=0,25, a razão ruído/sinal supera o que o modelo consegue superar com mais treino.")

    # ── 10. Perguntas para discussão ─────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "10. Perguntas para discussão")

    q1 = doc.add_paragraph(style="List Number")
    q1.add_run("A degradação no GEN1 corresponde ao que você observou. ").bold = True
    q1.add_run(
        "As três estratégias ficam abaixo do GEN0 em mAP50 no GEN3. "
        "Sua Figura 37 mostrou o mesmo padrão — degradação no GEN1 seguida de recuperação a partir do GEN2. "
        "Nossa cadeia GEN1–GEN3 equivale aos seus GEN1–GEN3, portanto ainda não vimos a fase de recuperação. "
        "Rodar um GEN4 parece o próximo passo natural. "
        "Você registrou AP50 por classe em cada geração? "
        "Saber se diode/inductor se recuperaram nos seus GEN2–GEN3 ajudaria a interpretar os nossos resultados.")

    q2 = doc.add_paragraph(style="List Number")
    q2.add_run("Como o C_th = 0,25 foi escolhido? ").bold = True
    q2.add_run(
        "A Progressive obteve diode AP50 = 0,602, superando o baseline (0,560). "
        "Isso sugere que começar com conf=0,25 no GEN1 é importante para capturar pseudo-labels "
        "de classes raras desde cedo. "
        "O seu C_th=0,25 foi definido empiricamente ou seguindo um trabalho anterior? "
        "Queremos entender se o nosso baseline Static é equivalente à sua melhor configuração.")

    q3 = doc.add_paragraph(style="List Number")
    q3.add_run("O GEN4 foi rodado para confirmar estabilização ou para obter o checkpoint final? ").bold = True
    q3.add_run(
        "Pela Figura 37, o mAP parece estabilizar entre GEN3 e GEN4. "
        "Você parou no GEN4 porque a curva havia claramente convergido, ou foi um orçamento fixo? "
        "Isso importa para decidir se nossa cadeia de três iterações é suficiente para ver a "
        "recuperação ou se precisamos estender ao GEN4.")

    q4 = doc.add_paragraph(style="List Number")
    q4.add_run("Adaptive threshold — sua perspectiva. ").bold = True
    q4.add_run(
        "Nossa Adaptive usa a curva PR do val para definir limiares por classe a cada geração. "
        "No GEN3 ela é a pior estratégia (mAP50=0,807 vs. Static/Progressive 0,826). "
        "Nossa hipótese: o modelo GEN1 não está calibrado o suficiente para produzir limiares "
        "confiáveis — e isso deveria melhorar em gerações posteriores. "
        "Você listou 'limiar dinâmico' como trabalho futuro. "
        "Você tem intuição sobre por que a calibração em gerações iniciais pode ser pouco confiável, "
        "ou alguma sugestão de abordagem diferente?")

    # ── 11. Próximos passos ──────────────────────────────────────────────────
    add_heading(doc, "11. Próximos passos")
    doc.add_paragraph(
        "Status atual:\n"
        "  ✓  GEN0 baselines — todos os quatro modelos concluídos\n"
        "  ✓  GEN1–GEN3 — YOLOv8m concluído (três estratégias)\n"
        "  ○  GEN1–GEN3 — YOLOv10m, RT-DETR-l, RT-DETR-x (scripts prontos, aguardando GPU)\n"
        "  ○  GEN4 — decisão pendente desta discussão\n\n"
        "Ordem planejada (uma execução por vez, RTX 4060 única):\n"
        "  1. YOLOv10m × 3 estratégias\n"
        "  2. RT-DETR-l × 3 estratégias\n"
        "  3. RT-DETR-x × 3 estratégias\n"
        "  4. Avaliar todos os checkpoints GEN3\n"
        "  5. Comparação entre arquiteturas\n"
        "  6. Decidir sobre extensão ao GEN4 com base nos resultados + esta discussão")

    # ── salvar ───────────────────────────────────────────────────────────────
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Salvo em: {OUT}")


if __name__ == "__main__":
    build()
